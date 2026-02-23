# D:\final\clinic_messages\views.py

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from django.urls import reverse_lazy, reverse
from django.db.models import Q
from django.contrib import messages
from django.utils import timezone
from rest_framework import generics
from rest_framework.permissions import IsAuthenticated
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status
from django.core.paginator import Paginator, EmptyPage, PageNotAnInteger
from django.forms import inlineformset_factory 
from django.db import transaction
from django.http import JsonResponse 
from django.conf import settings 
from django.core.files.storage import FileSystemStorage 
import os 
import uuid 
from django.http import HttpResponse
from docx import Document 
from docx.shared import Inches 
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from .models import Message, MessageRecipient, MessageAttachment, Notification # Ensure Notification is imported
from .forms import MessageForm, MessageAttachmentFormSet, MessageFilterForm 
from django import forms
from .serializers import MessageSerializer, MessageRecipientSerializer, UserSerializer
from django.contrib.auth import get_user_model
User = get_user_model()
from mammoth import convert_to_html
from .filters import MessageFilter 
from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import ListView
from django.utils.html import strip_tags
from django.db.models import Q, Value
from django.db.models.functions import Concat
from django.http import JsonResponse

# --- Helper function for mobile push notifications (if not defined elsewhere) ---
# This is a placeholder and should be implemented with a real push notification service
def send_mobile_push_notification(user, title, body, data=None):
    """
    This function sends a mobile push notification.
    You should implement this function with the actual logic of your push notification service (e.g., FCM).
    """
    print(f"Sending mobile push notification to {user.username}: Title='{title}', Body='{body}'")
    if data:
        print(f"Data: {data}")
    # Add actual code here to send notifications to mobile platforms (like FCM).


# --------------------------------------------------
# API Views
# --------------------------------------------------

class ConvertDocxToHtmlAPIView(APIView):
    """
    API endpoint to convert a DOCX file to HTML content.
    Expects a POST request with a 'docx_file' in request.FILES.
    """
    permission_classes = [IsAuthenticated] # Added permission for API access

    def post(self, request, *args, **kwargs):
        if 'docx_file' not in request.FILES:
            return Response({'error': 'No docx_file provided.'}, status=status.HTTP_400_BAD_REQUEST)

        docx_file = request.FILES['docx_file']

        if not docx_file.name.lower().endswith('.docx'):
            return Response({'error': 'Only .docx files are allowed.'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            result = convert_to_html(docx_file)
            html_content = result.value
            messages = result.messages

            return Response({'html_content': html_content, 'messages': [str(m) for m in messages]}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'error': f'Error converting docx to HTML: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)


class MessageMarkAsReadAPIView(APIView):
    """
    API endpoint to mark a specific message as read for the current user.
    """
    permission_classes = [IsAuthenticated]

    def post(self, request, pk, *args, **kwargs):
        try:
            message_recipient_status = get_object_or_404(MessageRecipient, message__pk=pk, recipient=request.user)
            if not message_recipient_status.is_read: # Use is_read field for consistency
                message_recipient_status.read_at = timezone.now()
                message_recipient_status.is_read = True
                message_recipient_status.save()
            return Response({'status': 'success', 'message': 'پیام خوانده شد.'}, status=status.HTTP_200_OK)
        except Exception as e:
            return Response({'status': 'error', 'message': str(e)}, status=status.HTTP_400_BAD_REQUEST)

class UserMessageListAPIView(generics.ListAPIView):
    """
    API endpoint to list messages received by the current user.
    """
    serializer_class = MessageRecipientSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return MessageRecipient.objects.filter(recipient=self.request.user).select_related('message', 'message__sender').order_by('-message__created_at')

class UnreadMessagesCountAPIView(APIView):
    """
    API endpoint to get the count of unread messages for the current user.
    """
    permission_classes = [IsAuthenticated]

    def get(self, request, *args, **kwargs):
        unread_count = MessageRecipient.objects.filter(
            recipient=request.user,
            is_read=False
        ).count()
        return Response({'unread_count': unread_count}, status=status.HTTP_200_OK)


class UserSearchAPIView(generics.ListAPIView):
    serializer_class = UserSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        query = self.request.query_params.get('q', '').strip()
        if not query:
            return User.objects.none()

        # ایجاد یک فیلد مجازی برای ترکیب نام و نام خانوادگی جهت جستجوی یکپارچه
        return User.objects.annotate(
            full_name=Concat('first_name', Value(' '), 'last_name')
        ).filter(
            Q(username__icontains=query) | 
            Q(first_name__icontains=query) | 
            Q(last_name__icontains=query) |
            Q(full_name__icontains=query) # جستجو در "نام + نام خانوادگی"
        ).distinct()[:20]

    def list(self, request, *args, **kwargs):
        queryset = self.get_queryset()
        results = []
        for user in queryset:
            # نمایش هوشمند: اگر نام داشت، نام کامل، در غیر این صورت نام کاربری
            full_name = f"{user.first_name} {user.last_name}".strip()
            display_text = full_name if full_name else user.username
            
            results.append({
                'id': user.id,
                'text': display_text
            })
        
        # خروجی استاندارد برای Select2 (استفاده از کلید 'items' طبق نیاز شما)
        return JsonResponse({'items': results, 'total_count': len(results)})


@login_required
def ckeditor_upload_file(request):
    """
    CKEditor file upload endpoint. Handles saving uploaded files.
    """
    if request.method == 'POST' and request.FILES.get('upload'):
        upload_file = request.FILES['upload']
        from django.core.files.storage import default_storage
        from django.core.files.base import ContentFile

        file_name = default_storage.save(os.path.join('ckeditor', upload_file.name), ContentFile(upload_file.read()))
        file_url = settings.MEDIA_URL + file_name

        response_data = {
            'url': file_url,
            'fileName': upload_file.name,
            'uploaded': 1
        }
        return JsonResponse(response_data)
    
    return JsonResponse({'uploaded': 0, 'error': {'message': 'Upload failed'}})


# --------------------------------------------------
# Views for Internal Messaging (HTML/Django Templates)
# --------------------------------------------------

@login_required
def export_composed_message_to_word(request):
    """
    Exports the content of a composed message (CKEditor body) to a downloadable Word file.
    """
    if request.method == 'POST':
        body_content_html = request.POST.get('body', '')
        subject = request.POST.get('subject', 'پیام جدید')
        
        document = Document()
        
        document.add_heading(f'موضوع: {subject}', level=1)
        document.add_paragraph(f'فرستنده: {request.user.get_full_name() or request.user.username}')
        document.add_paragraph(f'تاریخ: {timezone.now().strftime("%Y-%m-%d %H:%M")}')
        
        document.add_heading('متن پیام:', level=2)
        clean_body_text = strip_tags(body_content_html)
        document.add_paragraph(clean_body_text)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = f'attachment; filename="composed_message_{timezone.now().strftime("%Y%m%d%H%M%S")}.docx"'
        
        document.save(response)
        return response
    else:
        messages.error(request, "این عملیات فقط از طریق POST قابل انجام است.")
        return redirect(reverse_lazy('clinic_messages:message_create'))


@login_required
def export_message_to_word(request, pk):
    """
    Exports details of an existing message, including attachments, to a downloadable Word file.
    """
    message = get_object_or_404(Message, pk=pk)

    document = Document()

    document.add_heading(f'جزئیات پیام: {message.subject}', level=1)
    
    document.add_paragraph(f'فرستنده: {message.sender.get_full_name() or message.sender.username}')
    
    recipients_list = ", ".join([rec.recipient.get_full_name() or rec.recipient.username for rec in message.recipients_status.all()])
    document.add_paragraph(f'گیرندگان: {recipients_list}')

    document.add_paragraph(f'تاریخ ارسال: {message.created_at.strftime("%Y-%m-%d %H:%M")}')
    document.add_paragraph(f'موضوع: {message.subject}')
    
    document.add_heading('متن پیام:', level=2)
    clean_body = strip_tags(message.body)
    document.add_paragraph(clean_body)

    if message.attachments.exists():
        document.add_heading('پیوست‌ها:', level=2)
        for attachment in message.attachments.all():
            document.add_paragraph(f'- {attachment.file_name or attachment.file.name}')
            
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="message_{message.pk}.docx"'
    
    document.save(response)
    return response

@login_required(login_url=reverse_lazy('login'))
def message_inbox(request):
    """
    Displays the user's message inbox with filtering and pagination.
    """
    filter_form = MessageFilterForm(request.GET, request=request)
    
    received_messages_status = MessageRecipient.objects.filter(
        recipient=request.user
    ).select_related('message', 'message__sender').order_by('-message__created_at')

    message_filter = MessageFilter(request.GET, queryset=received_messages_status)
    filtered_queryset = message_filter.qs 

    # The filter_form.is_valid() check and manual query filtering are somewhat redundant
    # if MessageFilter handles all filtering. Keeping it as it was in the original code.
    if filter_form.is_valid():
        query = filter_form.cleaned_data.get('query')
        if query:
            filtered_queryset = filtered_queryset.filter(
                Q(message__subject__icontains=query) |
                Q(message__body__icontains=query) |
                Q(message__sender__username__icontains=query)
            )

    paginator = Paginator(filtered_queryset, 10) 
    page = request.GET.get('page')
    try:
        page_obj = paginator.page(page)
    except PageNotAnInteger:
        page_obj = paginator.page(1)
    except EmptyPage:
        page_obj = paginator.page(paginator.num_pages)

    context = {
        'page_title': 'صندوق ورودی پیام‌ها',
        'filter_form': filter_form,
        'page_obj': page_obj, 
        'unread_messages_count': MessageRecipient.objects.filter(recipient=request.user, is_read=False).count(),
    }
    return render(request, 'clinic_messages/message_inbox.html', context)


class SentMessageListView(LoginRequiredMixin, ListView):
    """
    Displays messages sent by the current user with filtering and pagination.
    """
    model = Message
    template_name = 'clinic_messages/message_sent.html'
    context_object_name = 'messages'
    paginate_by = 10

    def get_queryset(self):
        queryset = Message.objects.filter(sender=self.request.user).prefetch_related('recipients_status__recipient').order_by('-created_at')
        
        self.filter_form = MessageFilterForm(self.request.GET, request=self.request)
        if self.filter_form.is_valid():
            # Apply filters from the form
            query = self.filter_form.cleaned_data.get('query')
            user = self.filter_form.cleaned_data.get('user')
            start_date_fa = self.filter_form.cleaned_data.get('start_date_fa')
            end_date_fa = self.filter_form.cleaned_data.get('end_date_fa')

            if query:
                queryset = queryset.filter(Q(subject__icontains=query) | Q(body__icontains=query))
            
            if user:
                # Filter by recipient for sent messages
                queryset = queryset.filter(recipients_status__recipient=user).distinct() # Use distinct to avoid duplicate messages if multiple recipients match
            
            if start_date_fa:
                import jdatetime
                import datetime
                jdate = jdatetime.date.fromisoformat(start_date_fa.replace('/', '-'))
                gdate = jdate.togregorian()
                queryset = queryset.filter(created_at__gte=gdate)
            
            if end_date_fa:
                import jdatetime
                import datetime
                jdate = jdatetime.date.fromisoformat(end_date_fa.replace('/', '-'))
                gdate = jdate.togregorian()
                queryset = queryset.filter(created_at__lte=gdate + datetime.timedelta(days=1))

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['page_title'] = "پیام‌های ارسال شده"
        context['filter_form'] = self.filter_form
        
        messages = self.get_queryset()
        paginator = Paginator(messages, self.paginate_by)
        page = self.request.GET.get('page')
        try:
            page_obj = paginator.page(page)
        except PageNotAnInteger:
            page_obj = paginator.page(1)
        except EmptyPage:
            page_obj = paginator.page(paginator.num_pages)
        
        context['page_obj'] = page_obj
        return context


@login_required(login_url=reverse_lazy('login'))
def message_detail(request, pk):
    """
    Displays the details of a message and handles marking it as read
    and submitting a direct reply to the sender of the original message.
    """
    message = get_object_or_404(Message, pk=pk)
    
    message_recipient_status = None
    try:
        message_recipient_status = MessageRecipient.objects.get(message=message, recipient=request.user)
        if not message_recipient_status.is_read:
            message_recipient_status.is_read = True
            message_recipient_status.read_at = timezone.now()
            message_recipient_status.save()
            messages.info(request, "پیام به عنوان خوانده شده علامت‌گذاری شد.")
    except MessageRecipient.DoesNotExist:
        pass # Current user is not a recipient, or is the sender.

    initial_recipients = [message.sender.pk] 
    reply_subject = f"Re: {message.subject}"

    if request.method == 'POST':
        reply_form = MessageForm(request.POST, request.FILES) # request=request is not needed by MessageForm
        reply_attachment_formset = MessageAttachmentFormSet(request.POST, request.FILES, prefix='attachments_reply', queryset=MessageAttachment.objects.none())

        if reply_form.is_valid() and reply_attachment_formset.is_valid():
            try:
                with transaction.atomic():
                    new_message = reply_form.save(commit=False)
                    new_message.sender = request.user
                    new_message.parent_message = message
                    new_message.subject = reply_subject
                    new_message.save()
                    # No form.save_m2m() here as recipients field is disabled and handled manually below

                    # Only reply to the original sender
                    MessageRecipient.objects.create(message=new_message, recipient=message.sender)

                    for attachment_form in reply_attachment_formset:
                        if attachment_form.cleaned_data and not attachment_form.cleaned_data.get('DELETE', False):
                            attachment = attachment_form.save(commit=False)
                            attachment.message = new_message
                            attachment.save()
                    
                    messages.success(request, "پاسخ شما با موفقیت ارسال شد.")
                    return redirect('clinic_messages:message_detail', pk=message.pk)
            except Exception as e:
                messages.error(request, f"خطا در ارسال پاسخ: {e}")
        else:
            messages.error(request, "خطا در ارسال پاسخ. لطفا فرم را بررسی کنید.")
            print("Reply Form errors:", reply_form.errors)
            print("Reply Formset errors:", reply_attachment_formset.errors)
            # If form is invalid, re-render with errors
            context = {
                'page_title': "جزئیات پیام",
                'message': message,
                'reply_form': reply_form,
                'reply_attachment_formset': reply_attachment_formset,
            }
            return render(request, 'clinic_messages/message_detail.html', context)

    else: # GET request
        reply_form = MessageForm(
            initial={
                'recipients': [message.sender.pk],
                'subject': reply_subject,
                'parent_message': message.pk,
            }
        )
        reply_form.fields['recipients'].widget.attrs['disabled'] = 'disabled'
        reply_form.fields['subject'].widget.attrs['readonly'] = 'readonly'
        # Hide related fields for replies
        

        reply_attachment_formset = MessageAttachmentFormSet(prefix='attachments_reply', queryset=MessageAttachment.objects.none())
    
    context = {
        'page_title': "جزئیات پیام",
        'message': message,
        'reply_form': reply_form,
        'reply_attachment_formset': reply_attachment_formset,
    }
    return render(request, 'clinic_messages/message_detail.html', context)

@login_required(login_url=reverse_lazy('login'))
def message_create(request):
    """
    Handles creation of new messages with attachments.
    """
    if request.method == 'POST':
        form = MessageForm(request.POST, request.FILES) # Pass request.FILES for file uploads
        attachment_formset = MessageAttachmentFormSet(request.POST, request.FILES, prefix='attachments', instance=Message())
        
        if form.is_valid() and attachment_formset.is_valid():
            try:
                with transaction.atomic():
                    message = form.save(commit=False)
                    message.sender = request.user
                    message.save()
                    form.save_m2m() # For saving ManyToManyField (recipients)

                    for attachment_form in attachment_formset:
                        if attachment_form.cleaned_data and not attachment_form.cleaned_data.get('DELETE', False):
                            attachment = attachment_form.save(commit=False)
                            attachment.message = message
                            attachment.save()
                    
                    for recipient_user in form.cleaned_data['recipients']:
                        MessageRecipient.objects.create(
                            message=message,
                            recipient=recipient_user,
                            is_read=False
                        )
                    messages.success(request, "پیام با موفقیت ارسال شد.")
                    return redirect(reverse('clinic_messages:message_inbox'))
            except Exception as e:
                messages.error(request, f"خطا در ارسال پیام: {e}")
        else:
            messages.error(request, "لطفا خطاهای فرم را برطرف کنید.")
            print("Form errors:", form.errors)
            print("Formset errors:", attachment_formset.errors)

    else:
        form = MessageForm()
        attachment_formset = MessageAttachmentFormSet(prefix='attachments', instance=Message())

    context = {
        'page_title': 'ایجاد پیام جدید',
        'form': form,
        'attachment_formset': attachment_formset,
    }
    return render(request, 'clinic_messages/message_create.html', context)


@login_required(login_url=reverse_lazy('login'))
def message_reply(request, pk):
    """
    Handles replying to a message, sending the reply to the original sender and all original recipients.
    """
    parent_message = get_object_or_404(Message, pk=pk)

    if request.method == 'POST':
        form = MessageForm(request.POST, request.FILES)
        attachment_formset = MessageAttachmentFormSet(request.POST, request.FILES, prefix='attachments', instance=Message())
        
        if form.is_valid() and attachment_formset.is_valid():
            try:
                with transaction.atomic():
                    message = form.save(commit=False)
                    message.sender = request.user
                    message.parent_message = parent_message
                    message.subject = f"پاسخ: {parent_message.subject or 'بدون موضوع'}"
                    message.save()

                    for attachment_form in attachment_formset:
                        if attachment_form.cleaned_data and not attachment_form.cleaned_data.get('DELETE', False):
                            attachment = attachment_form.save(commit=False)
                            attachment.message = message
                            attachment.save()

                    recipients_for_reply = set()
                    recipients_for_reply.add(parent_message.sender) # Include original sender
                    
                    # Also include all other original recipients of the parent message
                    for mr in parent_message.recipients_status.all(): # Changed from received_messages to recipients_status for consistency
                        if mr.recipient != request.user:
                            recipients_for_reply.add(mr.recipient)

                    for recipient_user in recipients_for_reply:
                        MessageRecipient.objects.create(
                            message=message,
                            recipient=recipient_user,
                            is_read=False
                        )
                    
                    messages.success(request, "پاسخ شما با موفقیت ارسال شد.")
                    return redirect(reverse('clinic_messages:message_detail', args=[parent_message.pk]))
            except Exception as e:
                messages.error(request, f"خطا در ارسال پاسخ: {e}")
        else:
            messages.error(request, "لطفا خطاهای فرم را برطرف کنید.")
            print("Reply Form errors (message_reply):", form.errors)
            print("Reply Formset errors (message_reply):", attachment_formset.errors)
            # If form is invalid, re-render with errors
            context = {
                'page_title': f"پاسخ به پیام: {parent_message.subject or 'بدون موضوع'}",
                'form': form,
                'attachment_formset': attachment_formset,
                'parent_message': parent_message,
            }
            return render(request, 'clinic_messages/message_reply.html', context)

    else:
        form = MessageForm(initial={
            'subject': f"پاسخ: {parent_message.subject or 'بدون موضوع'}",
            'parent_message': parent_message.pk,
            'body': f"\n\n--- پیام اصلی از {parent_message.sender.get_full_name() or parent_message.sender.username} ---\n{parent_message.body}"
        })
        form.fields['recipients'].required = False
        form.fields['recipients'].widget.attrs['disabled'] = 'disabled'
        form.fields['subject'].widget.attrs['readonly'] = 'readonly'
        form.fields['related_visit'].widget = forms.HiddenInput()
        form.fields['related_drug_request'].widget = forms.HiddenInput()

        attachment_formset = MessageAttachmentFormSet(prefix='attachments', instance=Message())

    context = {
        'page_title': f"پاسخ به پیام: {parent_message.subject or 'بدون موضوع'}",
        'form': form,
        'attachment_formset': attachment_formset,
        'parent_message': parent_message,
    }
    return render(request, 'clinic_messages/message_reply.html', context)